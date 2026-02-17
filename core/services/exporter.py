import os
import json
import csv
from datetime import datetime

from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet

from docx import Document


# =====================================================
# MAIN EXPORT FUNCTION
# =====================================================

def export_interview(session, output_dir, format="pdf"):

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")

    fname = f"interview_{session.designation}_{ts}.{format}"
    path = os.path.join(output_dir, fname)

    answers = session.answers

    if format == "json":
        _export_json(path, session, answers)

    elif format == "csv":
        _export_csv(path, session, answers)

    elif format == "docx":
        _export_docx(path, session, answers)

    elif format == "pdf":
        _export_pdf(path, session, answers)

    else:
        raise ValueError(f"Unsupported export format: {format}")

    return path


# =====================================================
# JSON
# =====================================================

def _export_json(path, session, answers):

    data = {
        "candidate_name": session.candidate_name,
        "role": session.role_label,
        "designation": session.designation,
        "company": session.company,
        "answers": answers,
    }

    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2)


# =====================================================
# CSV
# =====================================================

def _export_csv(path, session, answers):

    with open(path, "w", newline="", encoding="utf-8") as f:

        writer = csv.writer(f)

        writer.writerow(["QuestionID", "Answer"])

        for qid, ans in answers.items():
            writer.writerow([qid, ans])


# =====================================================
# WORD
# =====================================================

def _export_docx(path, session, answers):

    doc = Document()

    doc.add_heading("Interview Report", 0)

    doc.add_paragraph(f"Candidate: {session.candidate_name}")
    doc.add_paragraph(f"Role: {session.role_label}")
    doc.add_paragraph(f"Company: {session.company}")

    doc.add_heading("Responses", level=1)

    for qid, ans in answers.items():
        doc.add_heading(qid, level=3)
        doc.add_paragraph(ans)

    doc.save(path)


# =====================================================
# PDF
# =====================================================

def _export_pdf(path, session, answers):

    styles = getSampleStyleSheet()

    story = []

    story.append(Paragraph("<b>Interview Report</b>", styles["Title"]))
    story.append(Spacer(1, 20))

    story.append(Paragraph(
        f"<b>Candidate:</b> {session.candidate_name}", styles["Normal"]
    ))

    story.append(Paragraph(
        f"<b>Role:</b> {session.role_label}", styles["Normal"]
    ))

    story.append(Paragraph(
        f"<b>Company:</b> {session.company}", styles["Normal"]
    ))

    story.append(Spacer(1, 20))

    table_data = [["Question ID", "Answer"]]

    for qid, ans in answers.items():
        table_data.append([qid, ans])

    table = Table(table_data, colWidths=[200, 300])

    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("GRID", (0, 0), (-1, -1), 1, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))

    story.append(table)

    doc = SimpleDocTemplate(path, pagesize=A4)
    doc.build(story)
